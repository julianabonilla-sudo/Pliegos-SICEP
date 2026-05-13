from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "templates", "plantilla_base.docx")
MAX_PROD_PLANTILLA = 6   # número de tablas de producto en la plantilla base
NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── Formateo de fechas ─────────────────────────────────────────────────────────

MESES = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
         "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]

def _fmt_fecha(iso: str) -> str:
    try:
        d = datetime.fromisoformat(iso.replace("T", " ").split(" ")[0])
        return f"{d.day} de {MESES[d.month]} de {d.year}"
    except Exception:
        return iso

def _fmt_fecha_hora(iso: str) -> str:
    try:
        dt = datetime.fromisoformat(iso.replace("T", " "))
        return f"{dt.day} de {MESES[dt.month]} de {dt.year} a las {dt.strftime('%H:%M')} horas"
    except Exception:
        return iso

def _mes_anio(iso: str) -> str:
    try:
        d = datetime.fromisoformat(iso.split("T")[0])
        return f"{MESES[d.month].capitalize()} {d.year}"
    except Exception:
        return iso

def _mes_anio_hoy() -> str:
    """Mes y año actual para la portada del documento."""
    now = datetime.now()
    return f"{MESES[now.month].capitalize()} {now.year}"

def _fmt_hora_del_dia(iso: str) -> str:
    """'2026-05-20T17:00' → '17:00 horas del 20 de mayo de 2026'"""
    try:
        dt = datetime.fromisoformat(iso.replace("T", " "))
        return f"{dt.strftime('%H:%M')} horas del {dt.day} de {MESES[dt.month]} de {dt.year}"
    except Exception:
        return iso


# ── Reemplazo de marcadores (nivel párrafo) ───────────────────────────────────

def _reemplazar_en_parrafos(parrafos, marcadores: dict):
    for p in parrafos:
        texto = "".join(r.text for r in p.runs)
        nuevo = texto
        for clave, valor in marcadores.items():
            if clave in nuevo:
                nuevo = nuevo.replace(clave, str(valor))
        if nuevo != texto and p.runs:
            p.runs[0].text = nuevo
            for r in p.runs[1:]:
                r.text = ""

def _reemplazar_en_doc(doc: Document, marcadores: dict):
    _reemplazar_en_parrafos(doc.paragraphs, marcadores)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                _reemplazar_en_parrafos(celda.paragraphs, marcadores)
    for seccion in doc.sections:
        _reemplazar_en_parrafos(seccion.header.paragraphs, marcadores)
        _reemplazar_en_parrafos(seccion.footer.paragraphs, marcadores)


# ── Reemplazo de marcadores a nivel XML (para clones de tablas) ───────────────

def _reemplazar_en_xml(elemento, marcadores: dict):
    """Reemplaza marcadores {{...}} en todos los nodos <w:t> del elemento XML."""
    for t_node in elemento.iter(f"{{{NS_W}}}t"):
        if t_node.text:
            for clave, valor in marcadores.items():
                if clave in t_node.text:
                    t_node.text = t_node.text.replace(clave, str(valor))


# ── Generación dinámica de tablas de productos ────────────────────────────────

def _marcadores_producto(prod: dict) -> dict:
    """Construye el dict de marcadores para un producto, usando la base {{prod_1_...}}."""
    garantias = (
        f"COMPRADOR: {prod.get('garantia_comprador', '')}\n"
        f"VENDEDOR: {prod.get('garantia_vendedor', '')}"
    )
    return {
        "{{prod_1_modalidad}}": prod["modalidad_suministro"],
        "{{prod_1_duracion}}":  (
            f"{_fmt_fecha(prod['duracion_inicio'])} "
            f"a {_fmt_fecha(prod['duracion_fin'])}"
        ),
        "{{prod_1_tamano}}":    f"{prod['tamano_mwh']:,.0f}",
        "{{prod_1_indexador}}": prod["indexador"],
        "{{prod_1_garantias}}": garantias,
    }

def _generar_tablas_productos(doc: Document, productos: list):
    """
    Reemplaza las N tablas de producto fijas de la plantilla por tablas dinámicas,
    una por cada producto de la convocatoria (funciona para cualquier N).

    Debe llamarse ANTES de _reemplazar_en_doc para que los marcadores {{prod_1_...}}
    aún estén presentes en el XML base.
    """
    tablas = list(doc.tables)
    n_template = min(MAX_PROD_PLANTILLA, len(tablas) - 1)

    if n_template == 0:
        return

    # Guardar copia del XML de la primera tabla de producto (tiene marcadores {{prod_1_...}})
    primera_prod_el = tablas[1]._element
    tabla_base_xml  = deepcopy(primera_prod_el)

    # Registrar el elemento anterior como punto de inserción
    elemento_ref = primera_prod_el.getprevious()
    padre        = primera_prod_el.getparent()

    # Eliminar TODAS las tablas de producto de la plantilla
    for i in range(n_template, 0, -1):
        if i < len(tablas):
            tablas[i]._element.getparent().remove(tablas[i]._element)

    # Insertar las N nuevas tablas en orden (insertamos de la última a la primera
    # usando addnext sobre el mismo elemento de referencia)
    for n in range(len(productos), 0, -1):
        prod = productos[n - 1]

        nueva_xml = deepcopy(tabla_base_xml)

        # Reemplazar marcadores {{prod_1_...}} con los datos de este producto
        marcadores_prod = _marcadores_producto(prod)
        # También actualizar el título de la tabla
        marcadores_prod["RESUMEN DEL PRODUCTO 1"] = f"RESUMEN DEL PRODUCTO {n}"
        _reemplazar_en_xml(nueva_xml, marcadores_prod)

        if elemento_ref is not None:
            elemento_ref.addnext(nueva_xml)
        else:
            padre.insert(0, nueva_xml)


# ── Sección de metodología libre ──────────────────────────────────────────────

def _insertar_parrafo_despues(ref_el, texto: str):
    """Inserta un párrafo simple con `texto` inmediatamente después de ref_el."""
    nuevo_p = OxmlElement("w:p")
    nuevo_r = OxmlElement("w:r")
    nuevo_t = OxmlElement("w:t")
    nuevo_t.text = texto
    nuevo_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    nuevo_r.append(nuevo_t)
    nuevo_p.append(nuevo_r)
    ref_el.addnext(nuevo_p)
    return nuevo_p

def _manejar_seccion_libre(doc: Document, productos: list):
    """
    Reemplaza {{metod_libre_seccion}} con la descripción de productos de metodología libre.
    Si no hay productos libres, limpia el placeholder.
    """
    libre_prods = [p for p in productos if p.get("metodologia_evaluacion") == "libre"]

    placeholder = None
    for p in doc.paragraphs:
        if "{{metod_libre_seccion}}" in p.text:
            placeholder = p
            break

    if placeholder is None:
        return

    if not libre_prods:
        # Limpiar el placeholder
        if placeholder.runs:
            placeholder.runs[0].text = ""
            for r in placeholder.runs[1:]:
                r.text = ""
        return

    # Hay productos libres: construir la sección
    nums = _fmt_lista_productos([p["numero"] for p in libre_prods])
    header = f"{nums}:"

    if placeholder.runs:
        placeholder.runs[0].text = header
        for r in placeholder.runs[1:]:
            r.text = ""
    else:
        placeholder.add_run(header)

    ultimo_el = placeholder._element
    for prod in libre_prods:
        desc = prod.get("descripcion_metodologia_libre") or (
            "Metodología definida por BIA ENERGY conforme al pliego de condiciones."
        )
        nuevo_p = _insertar_parrafo_despues(ultimo_el, desc)
        ultimo_el = nuevo_p


# ── Eliminación de secciones de metodología vacías ───────────────────────────

def _manejar_secciones_metodologia(doc: Document, productos: list):
    """
    Elimina del documento la sección de metodología mensual o anual completa
    cuando no hay productos que la usen.

    Debe ejecutarse ANTES de _reemplazar_en_doc porque necesita encontrar
    los párrafos por sus marcadores {{metod_mensual_header}} / {{metod_anual_header}}.

    La sección mensual abarca desde {{metod_mensual_header}} hasta (sin incluir)
    {{metod_anual_header}}.
    La sección anual abarca desde {{metod_anual_header}} hasta (sin incluir)
    {{metod_libre_seccion}}.
    """
    mensual = [p for p in productos if p.get("metodologia_evaluacion") == "mensual"]
    anual   = [p for p in productos if p.get("metodologia_evaluacion") == "anual"]

    if mensual and anual:
        return  # Ambas secciones tienen productos — nada que eliminar

    # Snapshot de párrafos (los índices se mantienen aunque eliminemos elementos)
    parrafos = list(doc.paragraphs)

    idx_mensual = idx_anual = idx_libre = None
    for i, p in enumerate(parrafos):
        t = p.text
        if idx_mensual is None and "{{metod_mensual_header}}" in t:
            idx_mensual = i
        elif idx_anual is None and "{{metod_anual_header}}" in t:
            idx_anual = i
        elif idx_libre is None and "{{metod_libre_seccion}}" in t:
            idx_libre = i

    # Eliminar sección mensual completa (header + descripción)
    if not mensual and idx_mensual is not None and idx_anual is not None:
        for p in parrafos[idx_mensual:idx_anual]:
            p._element.getparent().remove(p._element)

    # Eliminar sección anual completa (header + descripción)
    if not anual and idx_anual is not None and idx_libre is not None:
        for p in parrafos[idx_anual:idx_libre]:
            p._element.getparent().remove(p._element)


# ── Construcción de marcadores planos ─────────────────────────────────────────

def _fmt_lista_productos(numeros: list) -> str:
    """[1,2,3] → 'Productos 1, 2 y 3' | [1] → 'Producto 1'"""
    if not numeros:
        return ""
    nums = [str(n) for n in numeros]
    if len(nums) == 1:
        return f"Producto {nums[0]}"
    return "Productos " + ", ".join(nums[:-1]) + " y " + nums[-1]

def _construir_marcadores(data: dict) -> dict:
    cr = data["cronograma"]
    ct = data["contacto"]
    productos = data["productos"]

    # Agrupar productos por metodología para los headers
    mensual = [p["numero"] for p in productos if p.get("metodologia_evaluacion") == "mensual"]
    anual   = [p["numero"] for p in productos if p.get("metodologia_evaluacion") == "anual"]

    metod_mensual_header = (_fmt_lista_productos(mensual) + ":") if mensual else ""
    metod_anual_header   = (_fmt_lista_productos(anual)   + ":") if anual   else ""

    tipo_pliego = data.get("tipo_pliego", "definitivos").lower()

    return {
        # Tipo de pliego
        "{{tipo_pliego}}":               tipo_pliego,
        "{{tipo_pliego_upper}}":         tipo_pliego.upper(),
        "{{tipo_pliego_cap}}":           tipo_pliego.capitalize(),
        # Identificación
        "{{codigo_convocatoria}}":       data["codigo_convocatoria"],
        "{{periodo_inicio}}":            _fmt_fecha(data["periodo_contrato_inicio"]),
        "{{periodo_fin}}":               _fmt_fecha(data["periodo_contrato_fin"]),
        # Mes de indexación (único para toda la convocatoria)
        "{{mes_indexacion}}":            data.get("mes_indexacion", ""),
        # Portada: mes en que se crean los pliegos (fecha actual)
        "{{mes_publicacion}}":           _mes_anio_hoy(),
        # Contacto
        "{{contacto_nombre}}":           ct["nombre"],
        "{{contacto_telefono}}":         ct["telefono"],
        "{{contacto_email}}":            ct["email"],
        # Cronograma
        "{{fecha_publicacion_sicep}}":   _fmt_fecha(cr["publicacion_sicep"]),
        "{{fecha_limite_consultas}}":    _fmt_fecha_hora(cr["limite_consultas"]),
        "{{fecha_pliegos_definitivos}}": _fmt_fecha(cr["pliegos_definitivos"]),
        "{{fecha_limite_oferta}}":       _fmt_fecha_hora(cr["limite_oferta"]),
        "{{fecha_limite_oferta_hora}}":  _fmt_hora_del_dia(cr["limite_oferta"]),
        "{{fecha_habilitados_asic}}":    _fmt_fecha(cr["habilitados_asic"]),
        "{{fecha_audiencia_publica}}":   _fmt_fecha_hora(cr["audiencia_publica"]),
        "{{fecha_max_formalizacion}}":   _fmt_fecha(cr["max_formalizacion"]),
        "{{fecha_max_registro_asic}}":   _fmt_fecha(cr["max_registro_asic"]),
        # Headers de metodología (dinámicos según los productos)
        "{{metod_mensual_header}}":      metod_mensual_header,
        "{{metod_anual_header}}":        metod_anual_header,
    }


# ── Garantía de cumplimiento: pagaré vs pagaré+garantía ─────────────────────

def _manejar_garantia_cumplimiento(doc: Document, data: dict):
    """
    Si garantia_cumplimiento_tipo == 'pagare': elimina las secciones
    'Adicionalmente...' (garantía bancaria) en 8.1, 8.2 y minuta.
    Si es 'pagare_garantia': solo limpia los marcadores (mantiene el texto).
    """
    tipo = data.get("garantia_cumplimiento_tipo", "pagare_garantia")

    INI_FIN_PARES = [
        ("{{garantia_extra_vendedor_inicio}}",          "{{garantia_extra_vendedor_fin}}"),
        ("{{garantia_extra_comprador_inicio}}",         "{{garantia_extra_comprador_fin}}"),
        ("{{garantia_extra_vendedor_minuta_inicio}}",   "{{garantia_extra_vendedor_minuta_fin}}"),
        ("{{garantia_extra_comprador_minuta_inicio}}", "{{garantia_extra_comprador_minuta_fin}}"),
    ]

    parrafos = list(doc.paragraphs)

    for m_ini, m_fin in INI_FIN_PARES:
        idx_ini = idx_fin = None
        for i, p in enumerate(parrafos):
            if idx_ini is None and m_ini in p.text:
                idx_ini = i
            elif idx_ini is not None and idx_fin is None and m_fin in p.text:
                idx_fin = i

        if idx_ini is None or idx_fin is None:
            continue

        if tipo == "pagare":
            for p in parrafos[idx_ini:idx_fin + 1]:
                padre = p._element.getparent()
                if padre is not None:
                    padre.remove(p._element)
        else:
            for idx in (idx_ini, idx_fin):
                padre = parrafos[idx]._element.getparent()
                if padre is not None:
                    padre.remove(parrafos[idx]._element)


# ── Actualización de tabla de contenido (requiere Word instalado) ────────────

def _actualizar_toc(output_path: str):
    """Abre el documento en Word para actualizar TOC y números de página."""
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc_word = word.Documents.Open(os.path.abspath(output_path))
        doc_word.TablesOfContents(1).Update()
        doc_word.Save()
        doc_word.Close()
        word.Quit()
    except Exception:
        pass  # Word no disponible; el usuario puede presionar Ctrl+A, F9 al abrir


# ── Función principal ──────────────────────────────────────────────────────────

def generar_pliego(data: dict, output_path: str):
    productos = data["productos"]

    if os.path.exists(TEMPLATE_PATH):
        doc = Document(TEMPLATE_PATH)

        # 1. Generar tablas de productos dinámicamente
        #    (los marcadores {{prod_1_...}} deben estar intactos en el XML base)
        _generar_tablas_productos(doc, productos)

        # 2. Eliminar secciones de metodología que no aplican
        #    (debe ir ANTES del reemplazo para poder encontrar los marcadores)
        _manejar_secciones_metodologia(doc, productos)

        # 3. Construir marcadores planos y reemplazarlos en todo el documento
        marcadores = _construir_marcadores(data)
        _reemplazar_en_doc(doc, marcadores)

        # 4. Manejar la sección de metodología libre (inserta/limpia párrafos)
        _manejar_seccion_libre(doc, productos)

        # 5. Garantía de cumplimiento: pagaré solo o pagaré + garantía bancaria
        _manejar_garantia_cumplimiento(doc, data)

    else:
        marcadores = _construir_marcadores(data)
        doc = _crear_doc_desde_cero(data, marcadores)

    doc.save(output_path)
    _actualizar_toc(output_path)


# ── Fallback: documento desde cero (cuando no hay plantilla) ──────────────────

def _crear_doc_desde_cero(data: dict, marcadores: dict) -> Document:
    doc = Document()
    cr = data["cronograma"]
    ct = data["contacto"]

    titulo = doc.add_heading("PLIEGO DE CONDICIONES DEFINITIVOS", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Código de convocatoria: {data['codigo_convocatoria']}")
    doc.add_paragraph(
        f"Período a contratar: {marcadores['{{periodo_inicio}}']} "
        f"a {marcadores['{{periodo_fin}}']}"
    )

    doc.add_heading("1. Objeto", level=2)
    doc.add_paragraph(
        f"La presente convocatoria tiene por objeto: {data.get('objeto_contrato', '')}."
    )

    doc.add_heading("2. Funcionario encargado", level=2)
    doc.add_paragraph(f"Nombre: {ct['nombre']}")
    doc.add_paragraph(f"Teléfono: {ct['telefono']}")
    doc.add_paragraph(f"Correo: {ct['email']}")

    doc.add_heading("3. Cronograma del proceso", level=2)
    tabla_cr = doc.add_table(rows=1, cols=2)
    tabla_cr.style = "Table Grid"
    enc = tabla_cr.rows[0].cells
    enc[0].text = "Actividad"
    enc[1].text = "Fecha"
    for celda in enc:
        celda.paragraphs[0].runs[0].bold = True

    for actividad, clave in [
        ("Publicación aviso SICEP",        "{{fecha_publicacion_sicep}}"),
        ("Límite consultas a pliegos",      "{{fecha_limite_consultas}}"),
        ("Pliegos definitivos",             "{{fecha_pliegos_definitivos}}"),
        ("Límite entrega oferta",           "{{fecha_limite_oferta}}"),
        ("Remisión habilitados al ASIC",    "{{fecha_habilitados_asic}}"),
        ("Audiencia pública",               "{{fecha_audiencia_publica}}"),
        ("Máxima formalización de ofertas", "{{fecha_max_formalizacion}}"),
        ("Máximo registro ante ASIC",       "{{fecha_max_registro_asic}}"),
    ]:
        f = tabla_cr.add_row()
        f.cells[0].text = actividad
        f.cells[1].text = marcadores.get(clave, clave)

    doc.add_paragraph("")
    doc.add_heading("4. Garantía de seriedad de la oferta", level=2)
    doc.add_paragraph(data.get("tipo_garantia_seriedad", ""))

    doc.add_heading("5. Productos a contratar", level=2)
    for prod in productos:
        n = prod["numero"]
        doc.add_heading(f"Producto {n}", level=3)
        tabla = doc.add_table(rows=1, cols=2)
        tabla.style = "Table Grid"

        def _fila(campo, valor):
            fila = tabla.add_row()
            fila.cells[0].text = campo
            fila.cells[1].text = str(valor)
            fila.cells[0].paragraphs[0].runs[0].bold = True

        _fila("Modalidad de suministro", prod["modalidad_suministro"])
        _fila("Duración",
              f"{_fmt_fecha(prod['duracion_inicio'])} a {_fmt_fecha(prod['duracion_fin'])}")
        _fila("Tamaño (MWh)", f"{prod['tamano_mwh']:,.0f}")
        _fila("Indexador", prod["indexador"])
        met = prod.get("metodologia_evaluacion", "libre")
        _fila("Metodología", prod.get("descripcion_metodologia_libre", met))
        _fila("Garantía vendedor", prod.get("garantia_vendedor", ""))
        _fila("Garantía comprador", prod.get("garantia_comprador", ""))
        doc.add_paragraph("")

    return doc
