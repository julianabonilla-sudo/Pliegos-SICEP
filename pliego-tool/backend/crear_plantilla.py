"""
Script que genera plantilla_base.docx a partir del pliego de referencia.
Ejecutar UNA SOLA VEZ (o cada vez que el pliego de referencia cambie).

    python crear_plantilla.py
"""
import shutil, os, sys
from docx import Document

SRC  = os.path.join(os.path.dirname(__file__), "..", "..",
                    "PLIEGO DE CONDICIONES DEF CP-BIAC2026-003.docx")
DEST = os.path.join(os.path.dirname(__file__), "..", "templates", "plantilla_base.docx")

# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_celda(celda, texto: str):
    """Sobrescribe la celda con el texto dado; elimina párrafos extra si hay más de uno."""
    parrafos = celda.paragraphs
    for p in parrafos[1:]:
        p._element.getparent().remove(p._element)
    p0 = celda.paragraphs[0]
    if p0.runs:
        p0.runs[0].text = texto
        for r in p0.runs[1:]:
            r.text = ""
    else:
        p0.add_run(texto)

def _texto_parrafo_completo(p) -> str:
    return p.text  # python-docx ya recorre todos los w:t, incluidos los de w:hyperlink

def _reemplazar_en_parrafo(p, buscar: str, marcador: str) -> bool:
    """
    Reemplaza 'buscar' en el párrafo.
    Si el texto está en hipervínculos, los elimina y reconstruye el texto.
    """
    from docx.oxml.ns import qn
    texto_completo = _texto_parrafo_completo(p)
    if buscar not in texto_completo:
        return False
    nuevo_texto = texto_completo.replace(buscar, marcador)
    for hl in list(p._element.findall(qn("w:hyperlink"))):
        p._element.remove(hl)
    if p.runs:
        p.runs[0].text = nuevo_texto
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(nuevo_texto)
    return True

def _reemplazar_global(doc: Document, reemplazos: list):
    """Aplica una lista de (buscar, marcador) a todos los párrafos y celdas del doc."""
    objetos = list(doc.paragraphs)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                objetos.extend(celda.paragraphs)
    for seccion in doc.sections:
        objetos.extend(seccion.header.paragraphs)
        objetos.extend(seccion.footer.paragraphs)
    for p in objetos:
        for buscar, marcador in reemplazos:
            _reemplazar_en_parrafo(p, buscar, marcador)

# ── Reemplazos de texto en el cuerpo ──────────────────────────────────────────

REEMPLAZOS_TEXTO = [
    # Tipo de pliego (definitivos / preliminares)
    ("PLIEGO DE CONDICIONES DEFINITIVOS",            "PLIEGO DE CONDICIONES {{tipo_pliego_upper}}"),
    ("Pliegos de Condiciones definitivos",           "Pliegos de Condiciones {{tipo_pliego_cap}}"),
    # Código de convocatoria
    ("CP-BIAC2026-003",                              "{{codigo_convocatoria}}"),
    # Período global
    ("1 de enero de 2027 a 31 de diciembre de 2032", "{{periodo_inicio}} a {{periodo_fin}}"),
    # Mes de indexación (aparece en múltiples párrafos: precios, Anexo 2)
    ("marzo de 2026",                                "{{mes_indexacion}}"),
    # Contacto
    ("Nohora Mesa y Juliana Bonilla",                "{{contacto_nombre}}"),
    ("3112111595 - 3006301074",                      "{{contacto_telefono}}"),
    ("nohora@bia.app y juliana.bonilla@bia.app",     "{{contacto_email}}"),
    # Mes de publicación en portada (auto-calculado al generar)
    ("Mayo 2026",                                    "{{mes_publicacion}}"),
    # Fechas embebidas en párrafos
    ("el 2 de junio de 2026 a las 15:00 horas",      "el {{fecha_audiencia_publica}}"),
    ("hasta el día 23 de junio de 2026",             "hasta el día {{fecha_max_formalizacion}}"),
    # Sección 5.4: tres menciones de la fecha límite de oferta
    ("a más tardar el 20 de mayo de 2026",           "a más tardar el {{fecha_limite_oferta_dia}}"),
    ("a las 17:00 horas del 20 de mayo de 2026",     "a las {{fecha_limite_oferta_hora}}"),
    ("hasta las 17:00 horas del día 20 de mayo de 2026", "hasta {{fecha_limite_oferta}}"),
    # Headers de metodología (se vuelven dinámicos según los productos)
    ("Productos 1 y 2",                              "{{metod_mensual_header}}"),
    ("Productos 3, 4, 5 y 6:",                       "{{metod_anual_header}}"),
]

# ── Celdas de la tabla de cronograma (Tabla 0) ────────────────────────────────

CRONOGRAMA_CELDAS = {
    1: "{{fecha_publicacion_sicep}}",
    2: "{{fecha_limite_consultas}}",
    3: "{{fecha_pliegos_definitivos}}",
    4: "{{fecha_limite_oferta}}",
    5: "{{fecha_limite_oferta}}",
    6: "{{fecha_habilitados_asic}}",
    7: "{{fecha_audiencia_publica}}",
    8: "{{fecha_max_formalizacion}}",
    9: "{{fecha_max_registro_asic}}",
}

# ── Celdas variables de las tablas de productos (Tablas 1-6) ──────────────────
# Índice de fila (0-based) → campo
PROD_FILAS = {
    1: "modalidad",   # MODALIDAD DE SUMINISTRO
    3: "duracion",    # DURACIÓN DEL SUMINISTRO
    5: "tamano",      # TAMAÑO DEL NEGOCIO JURÍDICO (MWh)
    6: "indexador",   # INDEXADOR
    7: "garantias",   # GARANTÍAS
}

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    src  = os.path.abspath(SRC)
    dest = os.path.abspath(DEST)

    if not os.path.exists(src):
        print(f"ERROR: No se encontro el archivo fuente:\n  {src}", file=sys.stderr)
        sys.exit(1)

    os.makedirs(os.path.dirname(dest), exist_ok=True)
    shutil.copy2(src, dest)
    print(f"Copiado: {os.path.basename(src)} -> plantilla_base.docx")

    doc = Document(dest)

    # 1. Reemplazos de texto libre (incluye mes_indexacion y headers de metodología)
    _reemplazar_global(doc, REEMPLAZOS_TEXTO)
    print("[OK] Marcadores de texto aplicados")

    # 2. Tabla de cronograma (tabla 0)
    tabla_cr = doc.tables[0]
    for idx_fila, marcador in CRONOGRAMA_CELDAS.items():
        _set_celda(tabla_cr.rows[idx_fila].cells[1], marcador)
    print("[OK] Cronograma marcado")

    # 3. Tablas de productos (tablas 1-6) — solo columna de valor (col 1)
    for num_prod in range(1, 7):
        tabla = doc.tables[num_prod]
        for idx_fila, campo in PROD_FILAS.items():
            marcador = f"{{{{prod_{num_prod}_{campo}}}}}"
            _set_celda(tabla.rows[idx_fila].cells[1], marcador)
    print("[OK] Tablas de productos marcadas (1-6)")

    # 4. Insertar placeholder {{metod_libre_seccion}} en el párrafo vacío
    #    que sigue al final de la sección de metodología anual
    _insertar_placeholder_libre(doc)

    # 5. Insertar marcadores de garantía bancaria (pagaré vs pagaré+garantía)
    _insertar_marcadores_garantia(doc)

    doc.save(dest)
    print(f"\nPlantilla guardada en:\n  {dest}")


def _nueva_p_marcador(marcador: str):
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = marcador
    r.append(t)
    p.append(r)
    return p


def _insertar_marcadores_garantia(doc: Document):
    """
    Rodea las secciones 'Adicionalmente...' de garantía bancaria con marcadores
    {{garantia_extra_*_inicio}} / {{garantia_extra_*_fin}} para poder eliminarlas
    condicionalmente cuando el tipo de garantía sea solo pagaré.
    Se aplica a 4 bloques: vendedor pliego (8.1), comprador pliego (8.2),
    vendedor minuta y comprador minuta.
    """
    BLOQUES = [
        # (texto_inicio, texto_fin_boundary, marker_ini, marker_fin, ocurrencia)
        (
            "Adicionalmente, el VENDEDOR otorgar",
            "GARANTÍA DE CUMPLIMIENTO A CARGO DE BIA ENERGY",
            "{{garantia_extra_vendedor_inicio}}",
            "{{garantia_extra_vendedor_fin}}",
            0,
        ),
        (
            "Adicionalmente, BIA ENERGY otorgar",
            "AUTORIZACIÓN PARA CONSULTA DE DATOS",
            "{{garantia_extra_comprador_inicio}}",
            "{{garantia_extra_comprador_fin}}",
            0,
        ),
        (
            "Adicionalmente, el VENDEDOR otorgar",
            "Del COMPRADOR",
            "{{garantia_extra_vendedor_minuta_inicio}}",
            "{{garantia_extra_vendedor_minuta_fin}}",
            1,
        ),
        (
            "Adicionalmente, BIA ENERGY otorgar",
            "Incumplimiento y estimación anticipada de perjuicios",
            "{{garantia_extra_comprador_minuta_inicio}}",
            "{{garantia_extra_comprador_minuta_fin}}",
            1,
        ),
    ]

    parrafos = list(doc.paragraphs)

    for texto_ini, texto_fin, m_ini, m_fin, ocurrencia in BLOQUES:
        conteo = 0
        p_ini = p_fin = None

        for p in parrafos:
            if p_ini is None and texto_ini in p.text:
                if conteo == ocurrencia:
                    p_ini = p
                else:
                    conteo += 1
                continue
            if p_ini is not None and p_fin is None and texto_fin in p.text:
                p_fin = p
                break

        if p_ini is not None:
            p_ini._element.addprevious(_nueva_p_marcador(m_ini))
            print(f"[OK] {m_ini}")
        else:
            print(f"[WARN] No se encontró inicio para {m_ini}")

        if p_fin is not None:
            p_fin._element.addprevious(_nueva_p_marcador(m_fin))
            print(f"[OK] {m_fin}")
        else:
            print(f"[WARN] No se encontró fin para {m_fin}")


def _insertar_placeholder_libre(doc: Document):
    """
    Encuentra el párrafo vacío entre el final de la sección anual
    ('evaluará las ofertas de manera anual') y 'Regla para dirimir empates',
    y lo marca con {{metod_libre_seccion}}.
    """
    idx_anual_end = None
    for i, p in enumerate(doc.paragraphs):
        if "evaluar" in p.text and "manera anual" in p.text:
            idx_anual_end = i
            break

    if idx_anual_end is None:
        print("[WARN] No se encontro el fin de la seccion anual; placeholder libre no insertado")
        return

    for i in range(idx_anual_end + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if not p.text.strip():
            p.add_run("{{metod_libre_seccion}}")
            print("[OK] Placeholder {{metod_libre_seccion}} insertado en parrafo", i)
            return
        if "Regla para dirimir" in p.text:
            break

    print("[WARN] No se encontro parrafo vacio para {{metod_libre_seccion}}")


if __name__ == "__main__":
    main()
